from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List
import os

from ..database import get_db
from ..models import User, Document, DocumentChunk
from ..schemas import DocumentOut, DocumentChunkOut, DocumentChunkUpdate
from .auth import get_current_user
from ..services.doc_processor import DocProcessor
from ..services.vector_service import VectorService
from ..agent.graph import translation_agent
from ..agent.tools import format_glossary_for_prompt

router = APIRouter(prefix="/document", tags=["document"])

from concurrent.futures import ThreadPoolExecutor

def translate_single_chunk(chunk_id: int, user_id: int, db_session_factory, mode: str = "proposed"):
    """
    Translates a single chunk using its own DB session for thread safety.
    Cooperatively checks for document status changes to support cancellation.
    Supports 'proposed', 'baseline_a' (Gemini 1.5 Pro Zero-shot), and 'baseline_b' (GPT-4o Zero-shot).
    """
    db: Session = db_session_factory()
    try:
        chunk = db.query(DocumentChunk).filter(DocumentChunk.id == chunk_id).first()
        if not chunk or chunk.status == "done":
            return
            
        # Check if translation was cancelled
        document = db.query(Document).filter(Document.id == chunk.document_id).first()
        if not document or document.status != "processing":
            chunk.status = "pending"
            db.commit()
            return
            
        chunk.status = "translating"
        db.commit()
        
        try:
            if mode in ["baseline_a", "baseline_b"]:
                # Zero-shot baseline translation
                from ..llm_provider import get_llm
                from langchain_core.messages import SystemMessage, HumanMessage
                import re
                
                model_name = "google/gemini-1.5-pro" if mode == "baseline_a" else "openai/gpt-4o-mini"
                llm = get_llm(temperature=0.3, model_name=model_name)
                
                system_prompt = (
                    "You are a professional academic translator. Translate the given English sentence into Vietnamese directly. "
                    "Return ONLY the translated sentence, without any explanations, introductory text, or markdown code blocks."
                )
                user_prompt = chunk.original_text
                
                # Double-check cancellation right before calling LLM
                db.refresh(document)
                if document.status != "processing":
                    chunk.status = "pending"
                    db.commit()
                    return
                
                response = llm.invoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt)
                ])
                translated_text = response.content.strip()
                
                # Clean up markdown code blocks if the LLM wrapped the output in one
                translated_text = re.sub(r"^```(?:[a-zA-Z]+)?\n", "", translated_text, flags=re.IGNORECASE)
                translated_text = re.sub(r"\n```$", "", translated_text, flags=re.IGNORECASE)
                translated_text = translated_text.strip()
                
                chunk.translated_text = translated_text
                chunk.reviewer_feedback = f"Zero-shot baseline translation using {model_name}."
                chunk.status = "done"
                db.commit()
            else:
                # 1. Retrieve Glossary Terms using RAG
                matched_terms = VectorService.retrieve_glossary_for_text(
                    db=db,
                    user_id=user_id,
                    text=chunk.original_text
                )
                glossary_context_str = format_glossary_for_prompt(matched_terms)
                
                # 2. Retrieve past corrections (Few-shot feedback loop for LLM learning)
                past_corrections = VectorService.retrieve_past_corrections(
                    db=db,
                    user_id=user_id,
                    text=chunk.original_text
                )
                few_shot_lines = []
                for pc in past_corrections:
                    few_shot_lines.append(f"- English: '{pc['original_text']}'\n  Vietnamese Human Correction: '{pc['translated_text']}'")
                few_shot_context_str = "\n".join(few_shot_lines) if few_shot_lines else "None. Follow guidelines."
                
                # Double-check cancellation right before calling LLM
                db.refresh(document)
                if document.status != "processing":
                    chunk.status = "pending"
                    db.commit()
                    return
                    
                # 3. Prepare LangGraph Input State
                inputs = {
                    "original_text": chunk.original_text,
                    "context_window": chunk.context_window or chunk.original_text,
                    "glossary_context": glossary_context_str,
                    "few_shot_context": few_shot_context_str,
                    "translator_output": "",
                    "reviewer_output": {},
                    "review_attempts": 0,
                    "final_output": ""
                }
                
                # 3. Invoke LangGraph Graph
                try:
                    result = translation_agent.invoke(inputs)
                    chunk.translated_text = result.get("final_output", "")
                    rev_out = result.get("reviewer_output", {})
                    chunk.reviewer_feedback = rev_out.get("feedback", "")
                except Exception as graph_err:
                    print(f"Agentic graph failed for chunk {chunk_id}, falling back to offline NMT: {str(graph_err)}")
                    from ..services.nmt_service import NMTService
                    fallback_text = NMTService.translate(chunk.original_text)
                    chunk.translated_text = fallback_text
                    chunk.reviewer_feedback = f"Offline NMT Fallback (Agentic run failed: {str(graph_err)})"
                
                chunk.status = "done"
                db.commit()
            
        except Exception as chunk_err:
            print(f"Error translating chunk {chunk_id}: {str(chunk_err)}")
            chunk.status = "failed"
            chunk.reviewer_feedback = f"Translation failed: {str(chunk_err)}"
            db.commit()
            
    finally:
        db.close()

# Background Task to run the translation flow
def translate_document_background(document_id: int, user_id: int, db_session_factory, mode: str = "proposed"):
    db: Session = db_session_factory()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            return
            
        document.status = "processing"
        db.commit()
        
        # Get all chunks that are not translated yet
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id,
            DocumentChunk.status != "done"
        ).all()
        
        chunk_ids = [c.id for c in chunks]
        
        # Translate chunks in parallel (max 4 threads to avoid hitting API rate limits)
        if chunk_ids:
            with ThreadPoolExecutor(max_workers=4) as executor:
                # We use list to force completion of all generator items
                list(executor.map(
                    lambda cid: translate_single_chunk(cid, user_id, db_session_factory, mode),
                    chunk_ids
                ))
        
        # Refresh session to check status
        db.refresh(document)
        failed_count = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id, 
            DocumentChunk.status == "failed"
        ).count()
        
        if failed_count == len(chunks) and len(chunks) > 0:
            document.status = "failed"
        else:
            document.status = "translated"
        db.commit()
        
    except Exception as doc_err:
        print(f"Error in background translation for doc {document_id}: {str(doc_err)}")
        db.rollback()
    finally:
        db.close()

@router.post("/upload", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    file_bytes = file.file.read()
    file_name = file.filename
    
    # Process file and get chunks with context window
    try:
        processed_chunks = DocProcessor.process_file(file_name, file_bytes)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse document: {str(e)}"
        )
        
    if not processed_chunks:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document contains no readable text."
        )
        
    # Save the file to real uploads directory
    import uuid
    ext = file_name.split('.')[-1].lower()
    unique_filename = f"{uuid.uuid4()}.{ext}"
    uploads_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "uploads")
    os.makedirs(uploads_dir, exist_ok=True)
    saved_file_path = os.path.join(uploads_dir, unique_filename)
    with open(saved_file_path, "wb") as f:
        f.write(file_bytes)
        
    # Save document entry
    new_doc = Document(
        name=file_name,
        file_path=saved_file_path,
        file_type=ext,
        status="uploaded",
        user_id=current_user.id
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    # Save chunks
    for p_chunk in processed_chunks:
        db_chunk = DocumentChunk(
            document_id=new_doc.id,
            original_text=p_chunk["original_text"],
            context_window=p_chunk["context_window"],
            position_index=p_chunk["position_index"],
            status="pending"
        )
        db.add(db_chunk)
        
    db.commit()
    db.refresh(new_doc)
    return new_doc

@router.get("/", response_model=List[DocumentOut])
def get_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return db.query(Document).filter(Document.user_id == current_user.id).all()

@router.get("/{document_id}", response_model=DocumentOut)
def get_document_detail(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == current_user.id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    return document

@router.post("/{document_id}/translate", response_model=DocumentOut)
def translate_document(
    document_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == current_user.id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
        
    if document.status == "processing":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document is already being translated."
        )
        
    document.status = "processing"
    # Reset chunk statuses to pending if we are re-translating
    for chunk in document.chunks:
        if chunk.status != "done":
            chunk.status = "pending"
    db.commit()
    
    # We pass the SessionLocal class so the background thread can instantiate its own db session
    from ..database import SessionLocal
    background_tasks.add_task(
        translate_document_background,
        document_id=document.id,
        user_id=current_user.id,
        db_session_factory=SessionLocal
    )
    
    return document

@router.post("/{document_id}/stop", response_model=DocumentOut)
def stop_document_translation(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == current_user.id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
        
    if document.status != "processing":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document is not being translated."
        )
        
    # Set status back to uploaded to stop background threads
    document.status = "uploaded"
    db.commit()
    
    # Reset any actively translating chunks back to pending
    for chunk in document.chunks:
        if chunk.status == "translating":
            chunk.status = "pending"
    db.commit()
    
    db.refresh(document)
    return document

def check_and_update_glossary_from_feedback(db: Session, user_id: int, original_text: str, old_translation: str, new_translation: str):
    """
    Compares original text and translations to identify if the user changed a glossary term's translation.
    If so, updates the glossary database table automatically.
    """
    try:
        from ..models import Glossary
        from ..llm_provider import get_local_llm
        from langchain_core.messages import SystemMessage, HumanMessage
        import json
        import re
        
        # 1. Fetch user glossaries
        glossaries = db.query(Glossary).filter(Glossary.user_id == user_id).all()
        if not glossaries:
            return
            
        # Filter glossaries that actually appear in the original text
        applicable_glossaries = []
        for g in glossaries:
            if g.source_term.lower() in original_text.lower():
                applicable_glossaries.append({
                    "source_term": g.source_term,
                    "target_term": g.target_term
                })
                
        if not applicable_glossaries:
            return
            
        # 2. Call LLM to detect if any term was updated in the new translation
        llm = get_local_llm(temperature=0.1)
        system_prompt = (
            "You are a glossary manager. Analyze the original English sentence, the old translation, and the new human-corrected translation. "
            "Detect if the human changed the translation for any of these glossary terms: {glossary_list}.\n"
            "Return ONLY a JSON list of updated terms, for example: [{\"source_term\": \"AI Agent\", \"target_term\": \"Tác nhân AI mới\"}]. "
            "If no glossary term translation was changed, return an empty list: []."
        ).format(glossary_list=json.dumps(applicable_glossaries, ensure_ascii=False))
        
        user_prompt = (
            f"Original: {original_text}\n"
            f"Old Translation: {old_translation}\n"
            f"New Translation: {new_translation}"
        )
        
        response = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
        content = response.content.strip()
        
        # Clean JSON blocks
        cleaned = re.sub(r"^```(?:json)?", "", content, flags=re.IGNORECASE)
        cleaned = re.sub(r"```$", "", cleaned, flags=re.IGNORECASE)
        cleaned = cleaned.strip()
        
        updates = json.loads(cleaned)
        if updates and isinstance(updates, list):
            for update in updates:
                src = update.get("source_term")
                tgt = update.get("target_term")
                if src and tgt:
                    # Find and update in database
                    g_item = db.query(Glossary).filter(
                        Glossary.user_id == user_id,
                        Glossary.source_term == src
                    ).first()
                    if g_item and g_item.target_term != tgt:
                        print(f"Auto-updating glossary from feedback: '{src}' -> '{g_item.target_term}' updated to '{tgt}'")
                        g_item.target_term = tgt
                        db.commit()
    except Exception as e:
        print(f"Failed to auto-update glossary from feedback: {str(e)}")

@router.put("/chunk/{chunk_id}", response_model=DocumentChunkOut)
def update_chunk(
    chunk_id: int,
    chunk_data: DocumentChunkUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chunk = db.query(DocumentChunk).join(Document).filter(
        DocumentChunk.id == chunk_id,
        Document.user_id == current_user.id
    ).first()
    
    if not chunk:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document chunk not found"
        )
        
    old_translation = chunk.translated_text or ""
    new_translation = chunk_data.translated_text
    
    # Trigger glossary update from feedback in background/inline
    if old_translation != new_translation:
        check_and_update_glossary_from_feedback(
            db=db,
            user_id=current_user.id,
            original_text=chunk.original_text,
            old_translation=old_translation,
            new_translation=new_translation
        )
        
    chunk.translated_text = new_translation
    chunk.corrected_by_user = True  # Force true when edited by human for Translation Memory
    chunk.status = "done"
    db.commit()
    db.refresh(chunk)
    return chunk

from fastapi import Header
from typing import Optional
import jwt
from ..core.config import settings

@router.get("/{document_id}/export")
def export_translated_document(
    document_id: int,
    token: Optional[str] = None,
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    actual_token = token
    if not actual_token and authorization and authorization.startswith("Bearer "):
        actual_token = authorization.split(" ")[1]
        
    if not actual_token:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Mã xác thực token bị thiếu. Vui lòng đăng nhập lại."
        )
        
    try:
        payload = jwt.decode(actual_token, settings.JWT_SECRET, algorithms=[settings.JWT_ALGORITHM])
        user_id: int = payload.get("id")
        if user_id is None:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Token không hợp lệ.")
    except jwt.PyJWTError:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Token không hợp lệ hoặc đã hết hạn.")
        
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == user_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Không tìm thấy tài liệu tương ứng."
        )
        
    chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).order_by(DocumentChunk.position_index).all()
    
    from urllib.parse import quote
    
    # Check if the original document is a Word document and the file exists
    if document.file_type in ["docx", "doc"] and os.path.exists(document.file_path):
        from docx import Document as DocxDocument
        import io
        from fastapi.responses import StreamingResponse
        
        # Load the original document
        doc = DocxDocument(document.file_path)
        
        # Build dictionary of replacements: original_text -> translated_text
        replacements = {}
        for chunk in chunks:
            if chunk.translated_text:
                replacements[chunk.original_text.strip()] = chunk.translated_text.strip()
                
        # Helper to recursively get all paragraphs from body, tables, headers, footers
        def get_all_paragraphs(doc):
            paragraphs = list(doc.paragraphs)
            
            # Body tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
                        
            # Headers & Footers
            for section in doc.sections:
                if section.header:
                    paragraphs.extend(section.header.paragraphs)
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs.extend(cell.paragraphs)
                if section.footer:
                    paragraphs.extend(section.footer.paragraphs)
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs.extend(cell.paragraphs)
            return paragraphs
            
        paragraphs = get_all_paragraphs(doc)
        
        # Sort replacements by length descending to replace larger phrases/sentences first
        sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)
        
        for p in paragraphs:
            p_text = p.text
            if not p_text.strip():
                continue
                
            replaced = False
            for orig, trans in sorted_replacements:
                if orig in p_text:
                    p_text = p_text.replace(orig, trans)
                    replaced = True
                    
            if replaced:
                # Assign new text back to runs while preserving formatting of the first run
                if p.runs:
                    p.runs[0].text = p_text
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    p.text = p_text
                    
        # Save to memory stream
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        
        # Return DOCX file preserving the format, font, size, position, tables etc.
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        encoded_filename = quote(f"translated_{document.name}")
        return StreamingResponse(
            out_stream,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
        
    # Fallback to plain text for other file formats (e.g., pdf, txt) or if file path does not exist
    translated_texts = []
    for chunk in chunks:
        text = chunk.translated_text if chunk.translated_text else chunk.original_text
        translated_texts.append(text)
        
    full_text = " ".join(translated_texts)
    
    from fastapi.responses import Response
    encoded_filename = quote(f"translated_{document.name}.txt")
    return Response(
        content=full_text,
        media_type="text/plain",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )

@router.post("/chunk/{chunk_id}/retranslate", response_model=DocumentChunkOut)
def retranslate_chunk(
    chunk_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chunk = db.query(DocumentChunk).join(Document).filter(
        DocumentChunk.id == chunk_id,
        Document.user_id == current_user.id
    ).first()
    
    if not chunk:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Không tìm thấy phân đoạn văn bản."
        )
        
    chunk.status = "translating"
    db.commit()
    
    try:
        # 1. Retrieve Glossary Terms using RAG
        matched_terms = VectorService.retrieve_glossary_for_text(
            db=db,
            user_id=current_user.id,
            text=chunk.original_text
        )
        glossary_context_str = format_glossary_for_prompt(matched_terms)
        
        # 2. Retrieve past corrections (Human-in-the-loop dynamic fine-tuning context)
        past_corrections = VectorService.retrieve_past_corrections(
            db=db,
            user_id=current_user.id,
            text=chunk.original_text
        )
        few_shot_lines = []
        for pc in past_corrections:
            few_shot_lines.append(f"- English: '{pc['original_text']}'\n  Vietnamese Human Correction: '{pc['translated_text']}'")
        few_shot_context_str = "\n".join(few_shot_lines) if few_shot_lines else "None. Follow guidelines."
        
        # 3. Prepare LangGraph Input State
        inputs = {
            "original_text": chunk.original_text,
            "context_window": chunk.context_window or chunk.original_text,
            "glossary_context": glossary_context_str,
            "few_shot_context": few_shot_context_str,
            "translator_output": "",
            "reviewer_output": {},
            "review_attempts": 0,
            "final_output": ""
        }
        
        # 4. Invoke LangGraph Graph
        try:
            result = translation_agent.invoke(inputs)
            chunk.translated_text = result.get("final_output", "")
            rev_out = result.get("reviewer_output", {})
            chunk.reviewer_feedback = rev_out.get("feedback", "")
        except Exception as graph_err:
            print(f"Agentic graph failed for retranslation of chunk {chunk_id}, falling back to offline NMT: {str(graph_err)}")
            from ..services.nmt_service import NMTService
            fallback_text = NMTService.translate(chunk.original_text)
            chunk.translated_text = fallback_text
            chunk.reviewer_feedback = f"Offline NMT Fallback (Agentic run failed: {str(graph_err)})"
            
        chunk.status = "done"
        db.commit()
        db.refresh(chunk)
        return chunk
        
    except Exception as e:
        chunk.status = "failed"
        chunk.reviewer_feedback = f"Retranslation failed: {str(e)}"
        db.commit()
        db.refresh(chunk)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Dịch lại thất bại: {str(e)}"
        )
